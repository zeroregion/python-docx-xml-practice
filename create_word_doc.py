from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def create_document(body_pages=3):
    # 创建新文档
    doc = Document()

    # === 1. 添加封面（第一节） ===
    doc.add_heading('测试报告', 0).alignment = 1  # 居中对齐
    doc.add_paragraph('部门：技术研发部\n日期：2025年5月').alignment = 1  # 居中对齐
    doc.add_section(WD_SECTION.NEW_PAGE)  # 分节符（下一页）

    # === 2. 设置正文（第二节）页眉页脚 ===
    body_section = doc.sections[1]
    # 断开与封面节的链接
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False

    # 添加页眉
    header = body_section.header.paragraphs[0]
    header.text = "测试报告"
    header.alignment = 1 # 居中对齐

    # 添加带页码的页脚
    footer = body_section.footer.paragraphs[0]
    footer.text = "页码："
    _add_page_number(footer)  # 插入动态页码
    footer.alignment = 1  # 居中对齐

    # 设置正文页码从1开始
    _set_section_pagination(body_section, start=1)

    # === 3. 生成指定页数的正文内容 ===
    for page_num in range(body_pages):
        # 添加分页符（第一页不需要）
        if page_num > 0:
            doc.add_page_break()

        # 添加本页内容
        doc.add_heading(f'第{page_num + 1}节 核心内容', level=1)
        doc.add_paragraph(f"这是第 {page_num + 1} 页的详细技术文档内容..." + "此处填充示例文本。\n" * 10)

    # === 4. 添加封底（第三节） ===
    doc.add_section(WD_SECTION.NEW_PAGE)
    # 获取封底节并清除页眉页脚
    back_section = doc.sections[2]
    back_section.header.is_linked_to_previous = False
    back_section.footer.is_linked_to_previous = False

    # === 保存文档 ===
    doc.save("test.docx")


def _add_page_number(paragraph):
    """插入动态页码字段"""
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    run._r.append(instr_text)

    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)


def _set_section_pagination(section, start=1):
    """设置页码起始值"""
    sect_pr = section._sectPr
    pg_num_type = OxmlElement('w:pgNumType')
    pg_num_type.set(qn('w:start'), str(start))
    sect_pr.append(pg_num_type)


if __name__ == "__main__":
    create_document(body_pages=5)
